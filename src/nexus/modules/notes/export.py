"""Export functionality for notes to various formats."""

import io
import logging
from datetime import datetime
from pathlib import Path
from typing import List, Optional

from docx import Document
from docx.shared import Inches, Pt
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer
from weasyprint import HTML

from nexus.core.exceptions import ExportException

from .markdown_processor import MarkdownProcessor
from .models import Note

logger = logging.getLogger(__name__)


class NoteExporter:
    """Export notes to various formats."""

    def __init__(self):
        """Initialize exporter."""
        self.markdown_processor = MarkdownProcessor()

    def export_to_markdown(self, notes: List[Note], include_metadata: bool = True) -> str:
        """Export notes to Markdown format.

        Args:
            notes: List of notes to export
            include_metadata: Whether to include note metadata

        Returns:
            Markdown content
        """
        output = []

        for note in notes:
            output.append(f"# {note.title}\n")

            if include_metadata:
                output.append(f"**Created:** {note.created_at.strftime('%Y-%m-%d %H:%M')}\n")
                output.append(f"**Updated:** {note.updated_at.strftime('%Y-%m-%d %H:%M')}\n")

                if note.tags:
                    tags = ", ".join([tag.name for tag in note.tags])
                    output.append(f"**Tags:** {tags}\n")

                output.append("\n---\n\n")

            # Use markdown content if available, otherwise plain content
            content = note.content_markdown or note.content or ""
            output.append(content)
            output.append("\n\n")

        return "\n".join(output)

    def export_to_html(
        self,
        notes: List[Note],
        include_toc: bool = True,
        include_metadata: bool = True,
        custom_css: Optional[str] = None,
    ) -> str:
        """Export notes to HTML format.

        Args:
            notes: List of notes to export
            include_toc: Whether to include table of contents
            include_metadata: Whether to include note metadata
            custom_css: Optional custom CSS

        Returns:
            HTML content
        """
        default_css = """
        <style>
            body {
                font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, Oxygen, Ubuntu, Cantarell, sans-serif;
                line-height: 1.6;
                color: #333;
                max-width: 800px;
                margin: 0 auto;
                padding: 20px;
            }
            h1 { color: #2c3e50; border-bottom: 3px solid #3498db; padding-bottom: 10px; }
            h2 { color: #34495e; margin-top: 30px; }
            .metadata {
                background: #ecf0f1;
                padding: 10px;
                border-radius: 5px;
                margin-bottom: 20px;
                font-size: 0.9em;
            }
            .note {
                margin-bottom: 50px;
                page-break-after: always;
            }
            .tag {
                background: #3498db;
                color: white;
                padding: 3px 8px;
                border-radius: 3px;
                font-size: 0.8em;
                margin-right: 5px;
            }
            code {
                background: #f4f4f4;
                padding: 2px 6px;
                border-radius: 3px;
                font-family: 'Courier New', monospace;
            }
            pre {
                background: #2c3e50;
                color: #ecf0f1;
                padding: 15px;
                border-radius: 5px;
                overflow-x: auto;
            }
            table {
                border-collapse: collapse;
                width: 100%;
                margin: 20px 0;
            }
            th, td {
                border: 1px solid #ddd;
                padding: 12px;
                text-align: left;
            }
            th {
                background-color: #3498db;
                color: white;
            }
            .toc {
                background: #f8f9fa;
                padding: 20px;
                border-radius: 5px;
                margin-bottom: 30px;
            }
            .toc a {
                color: #3498db;
                text-decoration: none;
            }
            .toc a:hover {
                text-decoration: underline;
            }
        </style>
        """

        css = custom_css or default_css

        html = []
        html.append("<!DOCTYPE html>")
        html.append("<html>")
        html.append("<head>")
        html.append('<meta charset="UTF-8">')
        html.append("<title>Exported Notes</title>")
        html.append(css)
        html.append("</head>")
        html.append("<body>")

        # Table of contents
        if include_toc and len(notes) > 1:
            html.append('<div class="toc">')
            html.append("<h2>Table of Contents</h2>")
            html.append("<ul>")
            for note in notes:
                note_id = f"note-{note.id}"
                html.append(f'<li><a href="#{note_id}">{note.title}</a></li>')
            html.append("</ul>")
            html.append("</div>")

        # Notes content
        for note in notes:
            note_id = f"note-{note.id}"
            html.append(f'<div class="note" id="{note_id}">')
            html.append(f"<h1>{note.title}</h1>")

            if include_metadata:
                html.append('<div class="metadata">')
                html.append(f"<strong>Created:</strong> {note.created_at.strftime('%Y-%m-%d %H:%M')}<br>")
                html.append(f"<strong>Updated:</strong> {note.updated_at.strftime('%Y-%m-%d %H:%M')}<br>")

                if note.tags:
                    html.append("<strong>Tags:</strong> ")
                    for tag in note.tags:
                        html.append(f'<span class="tag">{tag.name}</span>')
                    html.append("<br>")

                html.append("</div>")

            # Convert content to HTML
            if note.content_html:
                html.append(note.content_html)
            elif note.content_markdown:
                content_html = self.markdown_processor.to_html(note.content_markdown)
                html.append(content_html)
            elif note.content:
                html.append(f"<p>{note.content}</p>")

            html.append("</div>")

        html.append("</body>")
        html.append("</html>")

        return "\n".join(html)

    def export_to_pdf(
        self, notes: List[Note], output_path: Optional[Path] = None
    ) -> bytes:
        """Export notes to PDF format using ReportLab.

        Args:
            notes: List of notes to export
            output_path: Optional output file path

        Returns:
            PDF content as bytes
        """
        try:
            # Create buffer
            buffer = io.BytesIO()

            # Create PDF document
            doc = SimpleDocTemplate(
                buffer,
                pagesize=letter,
                rightMargin=72,
                leftMargin=72,
                topMargin=72,
                bottomMargin=72,
            )

            # Container for elements
            elements = []

            # Styles
            styles = getSampleStyleSheet()
            title_style = ParagraphStyle(
                "CustomTitle",
                parent=styles["Heading1"],
                fontSize=24,
                textColor=colors.HexColor("#2c3e50"),
                spaceAfter=30,
            )
            heading_style = ParagraphStyle(
                "CustomHeading",
                parent=styles["Heading2"],
                fontSize=16,
                textColor=colors.HexColor("#34495e"),
                spaceAfter=12,
            )
            normal_style = styles["BodyText"]
            normal_style.fontSize = 11
            normal_style.leading = 14

            # Add notes
            for i, note in enumerate(notes):
                # Title
                elements.append(Paragraph(note.title, title_style))

                # Metadata
                metadata = f"Created: {note.created_at.strftime('%Y-%m-%d %H:%M')} | "
                metadata += f"Updated: {note.updated_at.strftime('%Y-%m-%d %H:%M')}"
                if note.tags:
                    tags = ", ".join([tag.name for tag in note.tags])
                    metadata += f" | Tags: {tags}"

                elements.append(Paragraph(metadata, styles["Italic"]))
                elements.append(Spacer(1, 0.2 * inch))

                # Content
                content = note.content or ""
                # Split content into paragraphs
                for para in content.split("\n\n"):
                    if para.strip():
                        elements.append(Paragraph(para, normal_style))
                        elements.append(Spacer(1, 0.1 * inch))

                # Page break between notes (except last)
                if i < len(notes) - 1:
                    elements.append(PageBreak())

            # Build PDF
            doc.build(elements)

            # Get PDF content
            pdf_content = buffer.getvalue()
            buffer.close()

            # Save to file if path provided
            if output_path:
                output_path.write_bytes(pdf_content)

            return pdf_content

        except Exception as e:
            logger.error(f"Failed to export to PDF: {e}")
            raise ExportException(f"PDF export failed: {str(e)}")

    def export_to_pdf_via_html(
        self, notes: List[Note], output_path: Optional[Path] = None
    ) -> bytes:
        """Export notes to PDF via HTML/CSS (better formatting).

        Args:
            notes: List of notes to export
            output_path: Optional output file path

        Returns:
            PDF content as bytes
        """
        try:
            # Generate HTML
            html_content = self.export_to_html(notes, include_toc=True, include_metadata=True)

            # Convert HTML to PDF using WeasyPrint
            pdf_bytes = HTML(string=html_content).write_pdf()

            # Save to file if path provided
            if output_path:
                output_path.write_bytes(pdf_bytes)

            return pdf_bytes

        except Exception as e:
            logger.error(f"Failed to export to PDF via HTML: {e}")
            raise ExportException(f"PDF export failed: {str(e)}")

    def export_to_word(
        self, notes: List[Note], output_path: Optional[Path] = None
    ) -> bytes:
        """Export notes to Microsoft Word format.

        Args:
            notes: List of notes to export
            output_path: Optional output file path

        Returns:
            Word document as bytes
        """
        try:
            # Create document
            doc = Document()

            # Set document properties
            doc.core_properties.title = "Exported Notes"
            doc.core_properties.author = "Nexus Platform"
            doc.core_properties.created = datetime.utcnow()

            for i, note in enumerate(notes):
                # Add title
                title = doc.add_heading(note.title, level=1)
                title_format = title.runs[0].font
                title_format.color.rgb = colors.HexColor("#2c3e50")

                # Add metadata
                metadata = doc.add_paragraph()
                metadata.add_run(f"Created: {note.created_at.strftime('%Y-%m-%d %H:%M')}\n").italic = True
                metadata.add_run(f"Updated: {note.updated_at.strftime('%Y-%m-%d %H:%M')}\n").italic = True

                if note.tags:
                    tags = ", ".join([tag.name for tag in note.tags])
                    tag_run = metadata.add_run(f"Tags: {tags}\n")
                    tag_run.italic = True
                    tag_run.font.color.rgb = colors.HexColor("#3498db")

                # Add content
                content = note.content or ""
                for para_text in content.split("\n\n"):
                    if para_text.strip():
                        doc.add_paragraph(para_text)

                # Page break between notes (except last)
                if i < len(notes) - 1:
                    doc.add_page_break()

            # Save to buffer
            buffer = io.BytesIO()
            doc.save(buffer)
            doc_content = buffer.getvalue()
            buffer.close()

            # Save to file if path provided
            if output_path:
                output_path.write_bytes(doc_content)

            return doc_content

        except Exception as e:
            logger.error(f"Failed to export to Word: {e}")
            raise ExportException(f"Word export failed: {str(e)}")

    def export_notes(
        self,
        notes: List[Note],
        format: str,
        output_path: Optional[Path] = None,
        **options,
    ) -> bytes:
        """Export notes to specified format.

        Args:
            notes: List of notes to export
            format: Export format (markdown, html, pdf, word)
            output_path: Optional output file path
            **options: Format-specific options

        Returns:
            Exported content as bytes

        Raises:
            ExportException: If export fails
        """
        if not notes:
            raise ExportException("No notes to export")

        format = format.lower()

        if format == "markdown":
            content = self.export_to_markdown(notes, **options)
            result = content.encode("utf-8")
        elif format == "html":
            content = self.export_to_html(notes, **options)
            result = content.encode("utf-8")
        elif format == "pdf":
            result = self.export_to_pdf_via_html(notes, output_path)
        elif format == "word":
            result = self.export_to_word(notes, output_path)
        else:
            raise ExportException(f"Unsupported export format: {format}")

        if output_path and format in ["markdown", "html"]:
            output_path.write_bytes(result)

        return result
