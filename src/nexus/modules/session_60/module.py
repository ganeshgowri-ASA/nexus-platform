"""
Session 60: Translation Module
Features: 60+ languages, document translation, real-time
"""
import asyncio
from typing import Any, Dict, List, Optional
from deep_translator import GoogleTranslator, MicrosoftTranslator
from langdetect import detect, detect_langs
import json
from loguru import logger
from docx import Document
import openpyxl

from ..base_module import BaseModule, ModuleConfig
from ...core.claude_client import ClaudeClient


class TranslationModule(BaseModule):
    """Multi-language translation with document support"""

    SUPPORTED_LANGUAGES = {
        "af": "Afrikaans", "sq": "Albanian", "am": "Amharic", "ar": "Arabic", "hy": "Armenian",
        "az": "Azerbaijani", "eu": "Basque", "be": "Belarusian", "bn": "Bengali", "bs": "Bosnian",
        "bg": "Bulgarian", "ca": "Catalan", "ceb": "Cebuano", "zh-CN": "Chinese (Simplified)",
        "zh-TW": "Chinese (Traditional)", "co": "Corsican", "hr": "Croatian", "cs": "Czech",
        "da": "Danish", "nl": "Dutch", "en": "English", "eo": "Esperanto", "et": "Estonian",
        "fi": "Finnish", "fr": "French", "fy": "Frisian", "gl": "Galician", "ka": "Georgian",
        "de": "German", "el": "Greek", "gu": "Gujarati", "ht": "Haitian Creole", "ha": "Hausa",
        "haw": "Hawaiian", "he": "Hebrew", "hi": "Hindi", "hmn": "Hmong", "hu": "Hungarian",
        "is": "Icelandic", "ig": "Igbo", "id": "Indonesian", "ga": "Irish", "it": "Italian",
        "ja": "Japanese", "jv": "Javanese", "kn": "Kannada", "kk": "Kazakh", "km": "Khmer",
        "rw": "Kinyarwanda", "ko": "Korean", "ku": "Kurdish", "ky": "Kyrgyz", "lo": "Lao",
        "la": "Latin", "lv": "Latvian", "lt": "Lithuanian", "lb": "Luxembourgish", "mk": "Macedonian",
        "mg": "Malagasy", "ms": "Malay", "ml": "Malayalam", "mt": "Maltese", "mi": "Maori",
        "mr": "Marathi", "mn": "Mongolian", "my": "Myanmar", "ne": "Nepali", "no": "Norwegian",
        "ny": "Nyanja", "or": "Odia", "ps": "Pashto", "fa": "Persian", "pl": "Polish",
        "pt": "Portuguese", "pa": "Punjabi", "ro": "Romanian", "ru": "Russian", "sm": "Samoan",
        "gd": "Scots Gaelic", "sr": "Serbian", "st": "Sesotho", "sn": "Shona", "sd": "Sindhi",
        "si": "Sinhala", "sk": "Slovak", "sl": "Slovenian", "so": "Somali", "es": "Spanish",
        "su": "Sundanese", "sw": "Swahili", "sv": "Swedish", "tl": "Tagalog", "tg": "Tajik",
        "ta": "Tamil", "tt": "Tatar", "te": "Telugu", "th": "Thai", "tr": "Turkish", "tk": "Turkmen",
        "uk": "Ukrainian", "ur": "Urdu", "ug": "Uyghur", "uz": "Uzbek", "vi": "Vietnamese",
        "cy": "Welsh", "xh": "Xhosa", "yi": "Yiddish", "yo": "Yoruba", "zu": "Zulu"
    }

    def __init__(self, claude_client: ClaudeClient, **kwargs):
        config = ModuleConfig(
            session=60,
            name="Translation",
            icon="🌍",
            description="60+ languages, document translation, real-time",
            version="1.0.0",
            features=["multilingual", "document_translation", "realtime", "language_detection"]
        )
        super().__init__(config, claude_client, **kwargs)

    def validate_input(self, input_data: Dict[str, Any]) -> bool:
        """Validate input"""
        return "action" in input_data

    async def detect_language(self, text: str) -> Dict[str, Any]:
        """Detect language of text"""
        try:
            lang = detect(text)
            probabilities = detect_langs(text)

            return {
                "success": True,
                "detected_language": lang,
                "language_name": self.SUPPORTED_LANGUAGES.get(lang, "Unknown"),
                "probabilities": [
                    {"language": p.lang, "name": self.SUPPORTED_LANGUAGES.get(p.lang, "Unknown"), "probability": p.prob}
                    for p in probabilities
                ]
            }
        except Exception as e:
            return self.handle_error(e, "detect_language")

    async def translate_text(
        self,
        text: str,
        target_language: str,
        source_language: str = "auto",
        engine: str = "google"
    ) -> Dict[str, Any]:
        """Translate text"""
        try:
            # Auto-detect source if needed
            if source_language == "auto":
                detection = await self.detect_language(text)
                source_language = detection.get("detected_language", "en")

            # Translate
            if engine == "google":
                translator = GoogleTranslator(source=source_language, target=target_language)
                translated = translator.translate(text)
            else:
                # Use Claude for high-quality translation
                prompt = f"Translate this text from {source_language} to {target_language}. Preserve formatting and meaning:\n\n{text}"
                translated = await self.claude.agenerate(prompt)

            return {
                "success": True,
                "original_text": text,
                "translated_text": translated,
                "source_language": source_language,
                "target_language": target_language,
                "engine": engine
            }

        except Exception as e:
            return self.handle_error(e, "translate_text")

    async def translate_document(
        self,
        file_path: str,
        target_language: str,
        source_language: str = "auto"
    ) -> Dict[str, Any]:
        """Translate document (DOCX, TXT, etc.)"""
        try:
            from pathlib import Path
            path = Path(file_path)

            if path.suffix == ".docx":
                doc = Document(file_path)
                translated_paragraphs = []

                for para in doc.paragraphs:
                    if para.text.strip():
                        result = await self.translate_text(para.text, target_language, source_language)
                        translated_paragraphs.append(result["translated_text"])
                    else:
                        translated_paragraphs.append("")

                # Create new document
                new_doc = Document()
                for text in translated_paragraphs:
                    new_doc.add_paragraph(text)

                output_path = path.parent / f"{path.stem}_{target_language}.docx"
                new_doc.save(output_path)

                return {
                    "success": True,
                    "source_file": file_path,
                    "output_file": str(output_path),
                    "paragraphs_translated": len(translated_paragraphs)
                }

            elif path.suffix == ".txt":
                text = path.read_text()
                result = await self.translate_text(text, target_language, source_language)
                output_path = path.parent / f"{path.stem}_{target_language}.txt"
                output_path.write_text(result["translated_text"])

                return {
                    "success": True,
                    "source_file": file_path,
                    "output_file": str(output_path)
                }

            else:
                return {"success": False, "error": f"Unsupported file type: {path.suffix}"}

        except Exception as e:
            return self.handle_error(e, "translate_document")

    async def batch_translate(
        self,
        texts: List[str],
        target_language: str,
        source_language: str = "auto"
    ) -> Dict[str, Any]:
        """Batch translate multiple texts"""
        try:
            results = []
            for text in texts:
                result = await self.translate_text(text, target_language, source_language)
                results.append(result)

            return {
                "success": True,
                "translations": results,
                "count": len(results)
            }

        except Exception as e:
            return self.handle_error(e, "batch_translate")

    def process(self, input_data: Dict[str, Any]) -> Dict[str, Any]:
        """Process translation request"""
        return asyncio.run(self.aprocess(input_data))

    async def aprocess(self, input_data: Dict[str, Any]) -> Dict[str, Any]:
        """Process translation request"""
        if not self.validate_input(input_data):
            return {"success": False, "error": "Invalid input"}

        action = input_data["action"]
        self.log_operation(action, input_data)

        try:
            if action == "detect":
                return await self.detect_language(input_data["text"])
            elif action == "translate":
                return await self.translate_text(
                    input_data["text"],
                    input_data["target_language"],
                    input_data.get("source_language", "auto"),
                    input_data.get("engine", "google")
                )
            elif action == "translate_document":
                return await self.translate_document(
                    input_data["file_path"],
                    input_data["target_language"],
                    input_data.get("source_language", "auto")
                )
            elif action == "batch":
                return await self.batch_translate(
                    input_data["texts"],
                    input_data["target_language"],
                    input_data.get("source_language", "auto")
                )
            elif action == "languages":
                return {
                    "success": True,
                    "languages": self.SUPPORTED_LANGUAGES
                }
            else:
                return {"success": False, "error": f"Unknown action: {action}"}

        except Exception as e:
            return self.handle_error(e, action)
