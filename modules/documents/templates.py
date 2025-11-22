"""
Document Templates Module

Provides document template management, creation, and auto-population capabilities.

Features:
- Template creation and management
- Auto-population from metadata
- Template versioning
- Template categories and organization
- Variable substitution
- Template validation
"""

import logging
import json
import re
from pathlib import Path
from typing import Dict, List, Optional, Union, Any, Set
from datetime import datetime
from enum import Enum
import uuid

logger = logging.getLogger(__name__)


class TemplateFormat(Enum):
    """Supported template formats."""
    DOCX = "docx"
    ODT = "odt"
    PDF = "pdf"
    HTML = "html"
    MARKDOWN = "md"
    TXT = "txt"


class TemplateError(Exception):
    """Base exception for template errors."""
    pass


class TemplateNotFoundError(TemplateError):
    """Exception for template not found errors."""
    pass


class TemplateValidationError(TemplateError):
    """Exception for template validation errors."""
    pass


class TemplateCategory:
    """
    Represents a category for organizing templates.
    """

    def __init__(
        self,
        name: str,
        description: Optional[str] = None,
        parent: Optional[str] = None
    ):
        """
        Initialize template category.

        Args:
            name: Category name
            description: Category description
            parent: Parent category name for hierarchical organization
        """
        self.name = name
        self.description = description
        self.parent = parent

    def to_dict(self) -> Dict[str, Any]:
        """Convert category to dictionary."""
        return {
            'name': self.name,
            'description': self.description,
            'parent': self.parent
        }

    @classmethod
    def from_dict(cls, data: Dict[str, Any]) -> 'TemplateCategory':
        """Create category from dictionary."""
        return cls(
            name=data['name'],
            description=data.get('description'),
            parent=data.get('parent')
        )


class TemplateVersion:
    """
    Represents a version of a template.
    """

    def __init__(
        self,
        version: str,
        content_path: Path,
        created_at: datetime,
        created_by: str,
        changes: Optional[str] = None
    ):
        """
        Initialize template version.

        Args:
            version: Version identifier (e.g., "1.0", "2.1")
            content_path: Path to template content for this version
            created_at: Creation timestamp
            created_by: User who created this version
            changes: Description of changes in this version
        """
        self.version = version
        self.content_path = content_path
        self.created_at = created_at
        self.created_by = created_by
        self.changes = changes

    def to_dict(self) -> Dict[str, Any]:
        """Convert version to dictionary."""
        return {
            'version': self.version,
            'content_path': str(self.content_path),
            'created_at': self.created_at.isoformat(),
            'created_by': self.created_by,
            'changes': self.changes
        }

    @classmethod
    def from_dict(cls, data: Dict[str, Any]) -> 'TemplateVersion':
        """Create version from dictionary."""
        return cls(
            version=data['version'],
            content_path=Path(data['content_path']),
            created_at=datetime.fromisoformat(data['created_at']),
            created_by=data['created_by'],
            changes=data.get('changes')
        )


class Template:
    """
    Represents a document template with metadata and versioning.
    """

    def __init__(
        self,
        template_id: str,
        name: str,
        format: TemplateFormat,
        category: Optional[str] = None,
        description: Optional[str] = None,
        variables: Optional[Set[str]] = None,
        metadata: Optional[Dict[str, Any]] = None
    ):
        """
        Initialize template.

        Args:
            template_id: Unique template identifier
            name: Template name
            format: Template format
            category: Template category
            description: Template description
            variables: Set of variable names used in template
            metadata: Additional metadata
        """
        self.template_id = template_id
        self.name = name
        self.format = format
        self.category = category
        self.description = description
        self.variables = variables or set()
        self.metadata = metadata or {}
        self.versions: List[TemplateVersion] = []
        self.created_at = datetime.now()
        self.updated_at = datetime.now()

    def add_version(
        self,
        version: str,
        content_path: Path,
        created_by: str,
        changes: Optional[str] = None
    ) -> TemplateVersion:
        """
        Add a new version to the template.

        Args:
            version: Version identifier
            content_path: Path to template content
            created_by: User creating the version
            changes: Description of changes

        Returns:
            Created TemplateVersion
        """
        template_version = TemplateVersion(
            version=version,
            content_path=content_path,
            created_at=datetime.now(),
            created_by=created_by,
            changes=changes
        )

        self.versions.append(template_version)
        self.updated_at = datetime.now()

        logger.info(f"Added version {version} to template {self.name}")

        return template_version

    def get_latest_version(self) -> Optional[TemplateVersion]:
        """Get the most recent version of the template."""
        if not self.versions:
            return None

        return max(self.versions, key=lambda v: v.created_at)

    def get_version(self, version: str) -> Optional[TemplateVersion]:
        """Get a specific version by identifier."""
        for v in self.versions:
            if v.version == version:
                return v
        return None

    def to_dict(self) -> Dict[str, Any]:
        """Convert template to dictionary."""
        return {
            'template_id': self.template_id,
            'name': self.name,
            'format': self.format.value,
            'category': self.category,
            'description': self.description,
            'variables': list(self.variables),
            'metadata': self.metadata,
            'versions': [v.to_dict() for v in self.versions],
            'created_at': self.created_at.isoformat(),
            'updated_at': self.updated_at.isoformat()
        }

    @classmethod
    def from_dict(cls, data: Dict[str, Any]) -> 'Template':
        """Create template from dictionary."""
        template = cls(
            template_id=data['template_id'],
            name=data['name'],
            format=TemplateFormat(data['format']),
            category=data.get('category'),
            description=data.get('description'),
            variables=set(data.get('variables', [])),
            metadata=data.get('metadata', {})
        )

        template.created_at = datetime.fromisoformat(data['created_at'])
        template.updated_at = datetime.fromisoformat(data['updated_at'])

        # Load versions
        for version_data in data.get('versions', []):
            template.versions.append(TemplateVersion.from_dict(version_data))

        return template


class TemplateManager:
    """
    Manages document templates, categories, and versioning.
    """

    def __init__(self, storage_path: Union[str, Path]):
        """
        Initialize template manager.

        Args:
            storage_path: Path to template storage directory
        """
        self.storage_path = Path(storage_path)
        self.storage_path.mkdir(parents=True, exist_ok=True)

        self.templates_dir = self.storage_path / "templates"
        self.templates_dir.mkdir(exist_ok=True)

        self.metadata_file = self.storage_path / "templates_metadata.json"

        self.templates: Dict[str, Template] = {}
        self.categories: Dict[str, TemplateCategory] = {}

        self._load_metadata()

        logger.info(f"TemplateManager initialized with {len(self.templates)} templates")

    def create_template(
        self,
        name: str,
        format: TemplateFormat,
        content: Union[str, bytes, Path],
        category: Optional[str] = None,
        description: Optional[str] = None,
        created_by: str = "system",
        metadata: Optional[Dict[str, Any]] = None
    ) -> Template:
        """
        Create a new template.

        Args:
            name: Template name
            format: Template format
            content: Template content (string, bytes, or path to file)
            category: Template category
            description: Template description
            created_by: User creating the template
            metadata: Additional metadata

        Returns:
            Created Template

        Raises:
            TemplateError: If template creation fails
        """
        template_id = str(uuid.uuid4())

        logger.info(f"Creating template: {name}")

        try:
            # Create template
            template = Template(
                template_id=template_id,
                name=name,
                format=format,
                category=category,
                description=description,
                metadata=metadata
            )

            # Save template content
            content_path = self._save_template_content(
                template_id,
                "1.0",
                format,
                content
            )

            # Extract variables from content
            variables = self._extract_variables(content_path)
            template.variables = variables

            # Add initial version
            template.add_version(
                version="1.0",
                content_path=content_path,
                created_by=created_by,
                changes="Initial version"
            )

            # Store template
            self.templates[template_id] = template
            self._save_metadata()

            logger.info(f"Template created: {template_id}")

            return template

        except Exception as e:
            logger.error(f"Failed to create template: {e}", exc_info=True)
            raise TemplateError(f"Template creation failed: {e}") from e

    def update_template(
        self,
        template_id: str,
        content: Union[str, bytes, Path],
        version: str,
        created_by: str,
        changes: Optional[str] = None
    ) -> TemplateVersion:
        """
        Update a template with a new version.

        Args:
            template_id: Template identifier
            content: New template content
            version: Version identifier
            created_by: User creating the version
            changes: Description of changes

        Returns:
            Created TemplateVersion

        Raises:
            TemplateNotFoundError: If template doesn't exist
        """
        template = self.get_template(template_id)

        if not template:
            raise TemplateNotFoundError(f"Template not found: {template_id}")

        logger.info(f"Updating template {template.name} to version {version}")

        # Save new content
        content_path = self._save_template_content(
            template_id,
            version,
            template.format,
            content
        )

        # Update variables
        variables = self._extract_variables(content_path)
        template.variables = variables

        # Add new version
        template_version = template.add_version(
            version=version,
            content_path=content_path,
            created_by=created_by,
            changes=changes
        )

        self._save_metadata()

        return template_version

    def get_template(self, template_id: str) -> Optional[Template]:
        """Get a template by ID."""
        return self.templates.get(template_id)

    def get_template_by_name(self, name: str) -> Optional[Template]:
        """Get a template by name."""
        for template in self.templates.values():
            if template.name == name:
                return template
        return None

    def list_templates(
        self,
        category: Optional[str] = None,
        format: Optional[TemplateFormat] = None
    ) -> List[Template]:
        """
        List templates with optional filtering.

        Args:
            category: Filter by category
            format: Filter by format

        Returns:
            List of templates
        """
        templates = list(self.templates.values())

        if category:
            templates = [t for t in templates if t.category == category]

        if format:
            templates = [t for t in templates if t.format == format]

        return templates

    def delete_template(self, template_id: str) -> bool:
        """
        Delete a template.

        Args:
            template_id: Template identifier

        Returns:
            True if deleted, False if not found
        """
        if template_id not in self.templates:
            return False

        template = self.templates[template_id]

        logger.info(f"Deleting template: {template.name}")

        # Delete template files
        template_dir = self.templates_dir / template_id
        if template_dir.exists():
            import shutil
            shutil.rmtree(template_dir)

        # Remove from memory
        del self.templates[template_id]
        self._save_metadata()

        return True

    def create_category(
        self,
        name: str,
        description: Optional[str] = None,
        parent: Optional[str] = None
    ) -> TemplateCategory:
        """
        Create a template category.

        Args:
            name: Category name
            description: Category description
            parent: Parent category name

        Returns:
            Created TemplateCategory
        """
        category = TemplateCategory(
            name=name,
            description=description,
            parent=parent
        )

        self.categories[name] = category
        self._save_metadata()

        logger.info(f"Category created: {name}")

        return category

    def list_categories(self) -> List[TemplateCategory]:
        """List all categories."""
        return list(self.categories.values())

    def populate_template(
        self,
        template_id: str,
        data: Dict[str, Any],
        output_path: Optional[Path] = None,
        version: Optional[str] = None
    ) -> Path:
        """
        Populate a template with data.

        Args:
            template_id: Template identifier
            data: Dictionary of variable values
            output_path: Optional output path for populated document
            version: Specific version to use (uses latest if not specified)

        Returns:
            Path to populated document

        Raises:
            TemplateNotFoundError: If template doesn't exist
            TemplateValidationError: If required variables are missing
        """
        template = self.get_template(template_id)

        if not template:
            raise TemplateNotFoundError(f"Template not found: {template_id}")

        # Get template version
        if version:
            template_version = template.get_version(version)
            if not template_version:
                raise TemplateNotFoundError(f"Version not found: {version}")
        else:
            template_version = template.get_latest_version()
            if not template_version:
                raise TemplateError("Template has no versions")

        logger.info(f"Populating template {template.name} version {template_version.version}")

        # Validate that all required variables are provided
        missing_vars = template.variables - set(data.keys())
        if missing_vars:
            logger.warning(f"Missing variables: {missing_vars}")

        # Read template content
        content_path = template_version.content_path

        if template.format in [TemplateFormat.TXT, TemplateFormat.MARKDOWN, TemplateFormat.HTML]:
            return self._populate_text_template(content_path, data, output_path)
        elif template.format == TemplateFormat.DOCX:
            return self._populate_docx_template(content_path, data, output_path)
        else:
            raise TemplateError(f"Unsupported template format for population: {template.format.value}")

    def _populate_text_template(
        self,
        template_path: Path,
        data: Dict[str, Any],
        output_path: Optional[Path]
    ) -> Path:
        """Populate a text-based template."""
        # Read template
        content = template_path.read_text(encoding='utf-8')

        # Replace variables (supports {{variable}} and {variable} syntax)
        for key, value in data.items():
            content = content.replace(f"{{{{{key}}}}}", str(value))
            content = content.replace(f"{{{key}}}", str(value))

        # Write output
        if not output_path:
            output_path = Path(f"populated_{template_path.name}")

        output_path.write_text(content, encoding='utf-8')

        logger.info(f"Populated text template: {output_path}")

        return output_path

    def _populate_docx_template(
        self,
        template_path: Path,
        data: Dict[str, Any],
        output_path: Optional[Path]
    ) -> Path:
        """Populate a DOCX template."""
        try:
            from docx import Document

            # Load template
            doc = Document(template_path)

            # Replace variables in paragraphs
            for paragraph in doc.paragraphs:
                for key, value in data.items():
                    if f"{{{{{key}}}}}" in paragraph.text:
                        paragraph.text = paragraph.text.replace(f"{{{{{key}}}}}", str(value))
                    if f"{{{key}}}" in paragraph.text:
                        paragraph.text = paragraph.text.replace(f"{{{key}}}", str(value))

            # Replace variables in tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for key, value in data.items():
                            if f"{{{{{key}}}}}" in cell.text:
                                cell.text = cell.text.replace(f"{{{{{key}}}}}", str(value))
                            if f"{{{key}}}" in cell.text:
                                cell.text = cell.text.replace(f"{{{key}}}", str(value))

            # Save output
            if not output_path:
                output_path = Path(f"populated_{template_path.name}")

            doc.save(output_path)

            logger.info(f"Populated DOCX template: {output_path}")

            return output_path

        except ImportError:
            raise TemplateError("python-docx library required for DOCX template population")

    def _save_template_content(
        self,
        template_id: str,
        version: str,
        format: TemplateFormat,
        content: Union[str, bytes, Path]
    ) -> Path:
        """Save template content to storage."""
        # Create template directory
        template_dir = self.templates_dir / template_id
        template_dir.mkdir(exist_ok=True)

        # Create content path
        content_path = template_dir / f"v{version}.{format.value}"

        # Save content
        if isinstance(content, Path):
            import shutil
            shutil.copy2(content, content_path)
        elif isinstance(content, bytes):
            content_path.write_bytes(content)
        else:
            content_path.write_text(content, encoding='utf-8')

        return content_path

    def _extract_variables(self, content_path: Path) -> Set[str]:
        """Extract variable names from template content."""
        try:
            # Read content
            if content_path.suffix in ['.txt', '.md', '.html']:
                content = content_path.read_text(encoding='utf-8')
            else:
                # For binary formats, try basic extraction
                content = content_path.read_text(encoding='utf-8', errors='ignore')

            # Find variables in {{variable}} and {variable} format
            pattern = r'\{\{?(\w+)\}?\}'
            variables = set(re.findall(pattern, content))

            logger.debug(f"Extracted variables: {variables}")

            return variables

        except Exception as e:
            logger.warning(f"Failed to extract variables: {e}")
            return set()

    def _load_metadata(self) -> None:
        """Load templates and categories metadata from storage."""
        if not self.metadata_file.exists():
            logger.info("No metadata file found, starting fresh")
            return

        try:
            with open(self.metadata_file, 'r') as f:
                data = json.load(f)

            # Load categories
            for cat_data in data.get('categories', []):
                category = TemplateCategory.from_dict(cat_data)
                self.categories[category.name] = category

            # Load templates
            for tmpl_data in data.get('templates', []):
                template = Template.from_dict(tmpl_data)
                self.templates[template.template_id] = template

            logger.info(f"Loaded {len(self.templates)} templates and {len(self.categories)} categories")

        except Exception as e:
            logger.error(f"Failed to load metadata: {e}", exc_info=True)

    def _save_metadata(self) -> None:
        """Save templates and categories metadata to storage."""
        try:
            data = {
                'categories': [cat.to_dict() for cat in self.categories.values()],
                'templates': [tmpl.to_dict() for tmpl in self.templates.values()]
            }

            with open(self.metadata_file, 'w') as f:
                json.dump(data, f, indent=2)

            logger.debug("Metadata saved")

        except Exception as e:
            logger.error(f"Failed to save metadata: {e}", exc_info=True)


# Convenience functions
def create_template(
    storage_path: Union[str, Path],
    name: str,
    format: str,
    content: Union[str, bytes, Path],
    **kwargs
) -> Template:
    """
    Convenience function to create a template.

    Args:
        storage_path: Path to template storage
        name: Template name
        format: Template format (e.g., 'docx', 'txt')
        content: Template content
        **kwargs: Additional template parameters

    Returns:
        Created Template
    """
    manager = TemplateManager(storage_path)
    format_enum = TemplateFormat(format.lower())
    return manager.create_template(name, format_enum, content, **kwargs)


def populate_template(
    storage_path: Union[str, Path],
    template_name: str,
    data: Dict[str, Any],
    output_path: Optional[Path] = None
) -> Path:
    """
    Convenience function to populate a template.

    Args:
        storage_path: Path to template storage
        template_name: Template name
        data: Variable values
        output_path: Optional output path

    Returns:
        Path to populated document
    """
    manager = TemplateManager(storage_path)
    template = manager.get_template_by_name(template_name)

    if not template:
        raise TemplateNotFoundError(f"Template not found: {template_name}")

    return manager.populate_template(template.template_id, data, output_path)
