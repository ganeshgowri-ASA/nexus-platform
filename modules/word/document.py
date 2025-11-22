"""
Document handling for Word Editor module.
"""
import json
import io
from datetime import datetime
from typing import Dict, List, Optional, Any
from pathlib import Path
from dataclasses import dataclass, asdict
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import markdown
import html2text
from bs4 import BeautifulSoup
from reportlab.lib.pagesizes import A4, letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from core.utils import (
    count_words,
    count_characters,
    estimate_reading_time,
    generate_document_id,
    get_timestamp,
)


@dataclass
class DocumentMetadata:
    """Document metadata."""
    id: str
    title: str
    created_at: str
    modified_at: str
    author: str = "Unknown"
    word_count: int = 0
    character_count: int = 0
    reading_time: Dict[str, int] = None
    version: int = 1
    tags: List[str] = None

    def __post_init__(self):
        if self.reading_time is None:
            self.reading_time = {"minutes": 0, "seconds": 0}
        if self.tags is None:
            self.tags = []


@dataclass
class Comment:
    """Document comment."""
    id: str
    text: str
    author: str
    timestamp: str
    position: int  # Character position in document
    resolved: bool = False


@dataclass
class Version:
    """Document version."""
    id: str
    content: str
    timestamp: str
    author: str
    comment: str = ""


class Document:
    """Document class for Word Editor."""

    def __init__(
        self,
        title: str = "Untitled Document",
        content: str = "",
        author: str = "Unknown",
    ):
        """
        Initialize document.

        Args:
            title: Document title
            content: Document content
            author: Document author
        """
        self.metadata = DocumentMetadata(
            id=generate_document_id(),
            title=title,
            created_at=get_timestamp(),
            modified_at=get_timestamp(),
            author=author,
        )
        self.content = content
        self.comments: List[Comment] = []
        self.versions: List[Version] = []
        self.formatting: Dict[str, Any] = {
            "font": "Arial",
            "font_size": 12,
            "color": "#000000",
            "bold": False,
            "italic": False,
            "underline": False,
            "alignment": "left",
        }

        # Update statistics
        self.update_statistics()

        # Save initial version
        self.save_version("Initial version")

    def update_content(self, content: str) -> None:
        """
        Update document content.

        Args:
            content: New content
        """
        self.content = content
        self.metadata.modified_at = get_timestamp()
        self.update_statistics()

    def update_statistics(self) -> None:
        """Update document statistics."""
        self.metadata.word_count = count_words(self.content)
        self.metadata.character_count = count_characters(self.content)
        self.metadata.reading_time = estimate_reading_time(self.content)

    def add_comment(
        self,
        text: str,
        author: str,
        position: int = 0,
    ) -> Comment:
        """
        Add a comment to the document.

        Args:
            text: Comment text
            author: Comment author
            position: Character position

        Returns:
            Created comment
        """
        comment = Comment(
            id=generate_document_id(),
            text=text,
            author=author,
            timestamp=get_timestamp(),
            position=position,
        )
        self.comments.append(comment)
        return comment

    def resolve_comment(self, comment_id: str) -> bool:
        """
        Mark a comment as resolved.

        Args:
            comment_id: Comment ID

        Returns:
            True if successful
        """
        for comment in self.comments:
            if comment.id == comment_id:
                comment.resolved = True
                return True
        return False

    def save_version(self, comment: str = "") -> Version:
        """
        Save a new version of the document.

        Args:
            comment: Version comment

        Returns:
            Created version
        """
        version = Version(
            id=generate_document_id(),
            content=self.content,
            timestamp=get_timestamp(),
            author=self.metadata.author,
            comment=comment,
        )
        self.versions.append(version)
        self.metadata.version = len(self.versions)
        return version

    def restore_version(self, version_id: str) -> bool:
        """
        Restore a previous version.

        Args:
            version_id: Version ID

        Returns:
            True if successful
        """
        for version in self.versions:
            if version.id == version_id:
                self.content = version.content
                self.metadata.modified_at = get_timestamp()
                self.update_statistics()
                self.save_version(f"Restored from version: {version_id}")
                return True
        return False

    def export_to_docx(self, file_path: Optional[Path] = None) -> bytes:
        """
        Export document to DOCX format.

        Args:
            file_path: Optional file path to save to

        Returns:
            DOCX file as bytes
        """
        doc = DocxDocument()

        # Add title
        title = doc.add_heading(self.metadata.title, 0)

        # Add content
        paragraphs = self.content.split("\n\n")
        for para_text in paragraphs:
            if para_text.strip():
                # Check for markdown headings
                if para_text.startswith("#"):
                    level = len(para_text) - len(para_text.lstrip("#"))
                    text = para_text.lstrip("#").strip()
                    doc.add_heading(text, level)
                else:
                    para = doc.add_paragraph(para_text)

        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        # Optionally save to file
        if file_path:
            doc.save(file_path)

        return buffer.getvalue()

    def export_to_pdf(self, file_path: Optional[Path] = None) -> bytes:
        """
        Export document to PDF format.

        Args:
            file_path: Optional file path to save to

        Returns:
            PDF file as bytes
        """
        buffer = io.BytesIO()
        pdf = SimpleDocTemplate(buffer, pagesize=A4)
        styles = getSampleStyleSheet()
        story = []

        # Add title
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            spaceAfter=30,
        )
        story.append(Paragraph(self.metadata.title, title_style))
        story.append(Spacer(1, 0.2 * inch))

        # Add content
        paragraphs = self.content.split("\n\n")
        for para_text in paragraphs:
            if para_text.strip():
                # Convert markdown-style formatting
                para_text = para_text.replace("**", "<b>").replace("**", "</b>")
                para_text = para_text.replace("*", "<i>").replace("*", "</i>")
                story.append(Paragraph(para_text, styles['BodyText']))
                story.append(Spacer(1, 0.1 * inch))

        pdf.build(story)
        buffer.seek(0)

        # Optionally save to file
        if file_path:
            with open(file_path, "wb") as f:
                f.write(buffer.getvalue())

        return buffer.getvalue()

    def export_to_html(self) -> str:
        """
        Export document to HTML format.

        Returns:
            HTML string
        """
        # Convert markdown to HTML
        html_content = markdown.markdown(
            self.content,
            extensions=['extra', 'codehilite', 'tables']
        )

        # Create full HTML document
        html = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{self.metadata.title}</title>
    <style>
        body {{
            font-family: {self.formatting['font']}, sans-serif;
            font-size: {self.formatting['font_size']}pt;
            line-height: 1.6;
            max-width: 800px;
            margin: 0 auto;
            padding: 20px;
        }}
        h1 {{ color: #333; }}
        h2 {{ color: #666; }}
        code {{ background-color: #f4f4f4; padding: 2px 5px; }}
        pre {{ background-color: #f4f4f4; padding: 10px; overflow-x: auto; }}
    </style>
</head>
<body>
    <h1>{self.metadata.title}</h1>
    {html_content}
</body>
</html>"""
        return html

    def export_to_markdown(self) -> str:
        """
        Export document to Markdown format.

        Returns:
            Markdown string
        """
        md = f"# {self.metadata.title}\n\n"
        md += f"*Author: {self.metadata.author}*\n\n"
        md += f"*Last modified: {self.metadata.modified_at}*\n\n"
        md += "---\n\n"
        md += self.content
        return md

    def import_from_docx(self, file_path: Path) -> None:
        """
        Import document from DOCX file.

        Args:
            file_path: Path to DOCX file
        """
        doc = DocxDocument(file_path)

        # Extract text
        content_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                content_parts.append(para.text)

        self.content = "\n\n".join(content_parts)
        self.metadata.modified_at = get_timestamp()
        self.update_statistics()

    def import_from_html(self, html_content: str) -> None:
        """
        Import document from HTML.

        Args:
            html_content: HTML string
        """
        # Convert HTML to text
        h = html2text.HTML2Text()
        h.ignore_links = False
        self.content = h.handle(html_content)
        self.metadata.modified_at = get_timestamp()
        self.update_statistics()

    def import_from_markdown(self, md_content: str) -> None:
        """
        Import document from Markdown.

        Args:
            md_content: Markdown string
        """
        self.content = md_content
        self.metadata.modified_at = get_timestamp()
        self.update_statistics()

    def to_dict(self) -> Dict[str, Any]:
        """
        Convert document to dictionary.

        Returns:
            Document as dictionary
        """
        return {
            "metadata": asdict(self.metadata),
            "content": self.content,
            "comments": [asdict(c) for c in self.comments],
            "versions": [asdict(v) for v in self.versions],
            "formatting": self.formatting,
        }

    def to_json(self) -> str:
        """
        Convert document to JSON.

        Returns:
            Document as JSON string
        """
        return json.dumps(self.to_dict(), indent=2)

    @classmethod
    def from_dict(cls, data: Dict[str, Any]) -> "Document":
        """
        Create document from dictionary.

        Args:
            data: Document data

        Returns:
            Document instance
        """
        doc = cls(
            title=data["metadata"]["title"],
            content=data["content"],
            author=data["metadata"]["author"],
        )
        doc.metadata = DocumentMetadata(**data["metadata"])
        doc.comments = [Comment(**c) for c in data.get("comments", [])]
        doc.versions = [Version(**v) for v in data.get("versions", [])]
        doc.formatting = data.get("formatting", doc.formatting)
        return doc

    @classmethod
    def from_json(cls, json_str: str) -> "Document":
        """
        Create document from JSON.

        Args:
            json_str: JSON string

        Returns:
            Document instance
        """
        data = json.loads(json_str)
        return cls.from_dict(data)
