"""
Export Module

Export OCR results to PDF, Word, Excel, JSON, and other formats.
"""

import logging
from typing import Dict, List, Any, Optional
from pathlib import Path
import json
from datetime import datetime

logger = logging.getLogger(__name__)


class ExportManager:
    """Main export manager for all formats"""

    def __init__(self):
        self.pdf_exporter = PDFExport()
        self.word_exporter = WordExport()
        self.excel_exporter = ExcelExport()
        self.json_exporter = JSONExport()
        self.logger = logging.getLogger(f"{__name__}.ExportManager")

    def export(
        self,
        data: Dict[str, Any],
        output_path: Path,
        format: str = "pdf",
        **kwargs
    ) -> bool:
        """
        Export OCR data to specified format

        Args:
            data: OCR result data
            output_path: Path to save output
            format: Export format
            **kwargs: Format-specific options

        Returns:
            Success status
        """
        try:
            format_lower = format.lower()

            if format_lower == "pdf":
                return self.pdf_exporter.export(data, output_path, **kwargs)
            elif format_lower in ["word", "docx"]:
                return self.word_exporter.export(data, output_path, **kwargs)
            elif format_lower in ["excel", "xlsx"]:
                return self.excel_exporter.export(data, output_path, **kwargs)
            elif format_lower == "json":
                return self.json_exporter.export(data, output_path, **kwargs)
            elif format_lower == "txt":
                return self._export_text(data, output_path)
            elif format_lower == "csv":
                return self._export_csv(data, output_path)
            else:
                self.logger.error(f"Unsupported format: {format}")
                return False

        except Exception as e:
            self.logger.error(f"Error exporting to {format}: {e}")
            return False

    def _export_text(self, data: Dict[str, Any], output_path: Path) -> bool:
        """Export as plain text"""
        try:
            text = data.get('text', '')
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(text)
            return True
        except Exception as e:
            self.logger.error(f"Error exporting text: {e}")
            return False

    def _export_csv(self, data: Dict[str, Any], output_path: Path) -> bool:
        """Export as CSV (for tables)"""
        try:
            import csv

            # Extract table data if available
            tables = data.get('tables', [])
            if not tables:
                return False

            with open(output_path, 'w', newline='', encoding='utf-8') as f:
                writer = csv.writer(f)
                for row in tables[0]:  # Export first table
                    writer.writerow(row)

            return True
        except Exception as e:
            self.logger.error(f"Error exporting CSV: {e}")
            return False


class PDFExport:
    """Export to PDF with searchable text layer"""

    def __init__(self):
        self.logger = logging.getLogger(f"{__name__}.PDFExport")

    def export(
        self,
        data: Dict[str, Any],
        output_path: Path,
        **kwargs
    ) -> bool:
        """
        Export to PDF with text layer

        Args:
            data: OCR result data
            output_path: Output PDF path
            **kwargs: Additional options

        Returns:
            Success status
        """
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.pdfgen import canvas
            from reportlab.lib.units import inch

            # Create PDF
            c = canvas.Canvas(str(output_path), pagesize=letter)
            width, height = letter

            # Add title
            c.setFont("Helvetica-Bold", 16)
            c.drawString(inch, height - inch, "OCR Result")

            # Add metadata
            c.setFont("Helvetica", 10)
            y_position = height - 1.5 * inch

            metadata = data.get('metadata', {})
            for key, value in metadata.items():
                c.drawString(inch, y_position, f"{key}: {value}")
                y_position -= 0.3 * inch

            # Add text content
            c.setFont("Helvetica", 12)
            y_position -= 0.5 * inch

            text = data.get('text', '')
            lines = text.split('\n')

            for line in lines:
                if y_position < inch:
                    c.showPage()
                    y_position = height - inch

                # Wrap long lines
                if len(line) > 80:
                    words = line.split()
                    current_line = ""
                    for word in words:
                        if len(current_line + word) < 80:
                            current_line += word + " "
                        else:
                            c.drawString(inch, y_position, current_line)
                            y_position -= 0.3 * inch
                            current_line = word + " "
                    if current_line:
                        c.drawString(inch, y_position, current_line)
                        y_position -= 0.3 * inch
                else:
                    c.drawString(inch, y_position, line)
                    y_position -= 0.3 * inch

            c.save()
            self.logger.info(f"Exported to PDF: {output_path}")
            return True

        except ImportError:
            self.logger.error("reportlab not installed")
            return False
        except Exception as e:
            self.logger.error(f"Error exporting to PDF: {e}")
            return False


class WordExport:
    """Export to Word document"""

    def __init__(self):
        self.logger = logging.getLogger(f"{__name__}.WordExport")

    def export(
        self,
        data: Dict[str, Any],
        output_path: Path,
        **kwargs
    ) -> bool:
        """
        Export to Word document

        Args:
            data: OCR result data
            output_path: Output Word path
            **kwargs: Additional options

        Returns:
            Success status
        """
        try:
            from docx import Document
            from docx.shared import Pt

            doc = Document()

            # Add title
            title = doc.add_heading('OCR Result', 0)

            # Add metadata
            metadata = data.get('metadata', {})
            if metadata:
                doc.add_heading('Metadata', 2)
                for key, value in metadata.items():
                    doc.add_paragraph(f"{key}: {value}")

            # Add text content
            doc.add_heading('Extracted Text', 2)

            text = data.get('text', '')
            paragraphs = text.split('\n\n')

            for para in paragraphs:
                if para.strip():
                    p = doc.add_paragraph(para.strip())
                    p.style.font.size = Pt(12)

            # Add tables if present
            tables = data.get('tables', [])
            if tables:
                doc.add_heading('Tables', 2)
                for table_data in tables:
                    if table_data:
                        table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                        table.style = 'Light Grid Accent 1'

                        for i, row in enumerate(table_data):
                            for j, cell_value in enumerate(row):
                                table.rows[i].cells[j].text = str(cell_value)

            doc.save(str(output_path))
            self.logger.info(f"Exported to Word: {output_path}")
            return True

        except ImportError:
            self.logger.error("python-docx not installed")
            return False
        except Exception as e:
            self.logger.error(f"Error exporting to Word: {e}")
            return False


class ExcelExport:
    """Export to Excel"""

    def __init__(self):
        self.logger = logging.getLogger(f"{__name__}.ExcelExport")

    def export(
        self,
        data: Dict[str, Any],
        output_path: Path,
        **kwargs
    ) -> bool:
        """
        Export to Excel

        Args:
            data: OCR result data
            output_path: Output Excel path
            **kwargs: Additional options

        Returns:
            Success status
        """
        try:
            import pandas as pd
            from openpyxl import Workbook
            from openpyxl.styles import Font, Alignment

            wb = Workbook()
            ws = wb.active
            ws.title = "OCR Result"

            row = 1

            # Add metadata
            metadata = data.get('metadata', {})
            if metadata:
                ws.cell(row, 1, "Metadata").font = Font(bold=True, size=14)
                row += 1

                for key, value in metadata.items():
                    ws.cell(row, 1, key).font = Font(bold=True)
                    ws.cell(row, 2, str(value))
                    row += 1
                row += 1

            # Add text content
            ws.cell(row, 1, "Extracted Text").font = Font(bold=True, size=14)
            row += 2

            text = data.get('text', '')
            lines = text.split('\n')
            for line in lines:
                ws.cell(row, 1, line)
                row += 1

            # Add tables in separate sheets
            tables = data.get('tables', [])
            for i, table_data in enumerate(tables):
                if table_data:
                    sheet_name = f"Table_{i + 1}"
                    ws_table = wb.create_sheet(title=sheet_name)

                    for row_idx, row_data in enumerate(table_data, 1):
                        for col_idx, cell_value in enumerate(row_data, 1):
                            ws_table.cell(row_idx, col_idx, str(cell_value))

            wb.save(str(output_path))
            self.logger.info(f"Exported to Excel: {output_path}")
            return True

        except ImportError:
            self.logger.error("openpyxl not installed")
            return False
        except Exception as e:
            self.logger.error(f"Error exporting to Excel: {e}")
            return False


class JSONExport:
    """Export to JSON"""

    def __init__(self):
        self.logger = logging.getLogger(f"{__name__}.JSONExport")

    def export(
        self,
        data: Dict[str, Any],
        output_path: Path,
        pretty: bool = True,
        **kwargs
    ) -> bool:
        """
        Export to JSON

        Args:
            data: OCR result data
            output_path: Output JSON path
            pretty: Pretty print JSON
            **kwargs: Additional options

        Returns:
            Success status
        """
        try:
            # Add timestamp
            export_data = {
                'timestamp': datetime.now().isoformat(),
                'data': data
            }

            with open(output_path, 'w', encoding='utf-8') as f:
                if pretty:
                    json.dump(export_data, f, indent=2, ensure_ascii=False)
                else:
                    json.dump(export_data, f, ensure_ascii=False)

            self.logger.info(f"Exported to JSON: {output_path}")
            return True

        except Exception as e:
            self.logger.error(f"Error exporting to JSON: {e}")
            return False

    def export_structured(
        self,
        data: Dict[str, Any],
        output_path: Path,
        **kwargs
    ) -> bool:
        """Export with structured format"""
        try:
            structured = {
                'metadata': data.get('metadata', {}),
                'text': data.get('text', ''),
                'confidence': data.get('confidence', 0.0),
                'pages': data.get('pages', []),
                'tables': data.get('tables', []),
                'forms': data.get('forms', {}),
                'entities': data.get('entities', {}),
            }

            return self.export(structured, output_path, **kwargs)

        except Exception as e:
            self.logger.error(f"Error exporting structured JSON: {e}")
            return False
